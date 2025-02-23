"""
Slide Explainer - Main Application
--------------------------------
A Streamlit application for interactive slide explanations using AI and TTS.
"""

import os
import asyncio
import streamlit as st
from pathlib import Path
from dotenv import load_dotenv
import re
from utils.audio_handler import AudioHandler
from utils.llm_interface import LLMInterface
from utils.slide_processor import SlideProcessor
import pandas as pd
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from utils.context_manager import ContextManager, DEFAULT_CONSTRAINTS
import markdown
from bs4 import BeautifulSoup
import docx

# Load environment variables
load_dotenv()

# Configure Streamlit page
st.set_page_config(
    page_title="Slide Explainer",
    page_icon="📊",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state
if "current_slide_index" not in st.session_state:
    st.session_state.current_slide_index = 0
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "cached_responses" not in st.session_state:
    st.session_state.cached_responses = {}
if "slides" not in st.session_state:
    st.session_state.slides = []
if "is_caching" not in st.session_state:
    st.session_state.is_caching = False
if "next_cache_index" not in st.session_state:
    st.session_state.next_cache_index = None
if "page_cache" not in st.session_state:
    st.session_state.page_cache = {}
if "notes_file" not in st.session_state:
    st.session_state.notes_file = None
if "saved_notes" not in st.session_state:
    st.session_state.saved_notes = set()
if "edited_explanations" not in st.session_state:
    st.session_state.edited_explanations = {}
if "notes_format" not in st.session_state:
    st.session_state.notes_format = "docx"
if "context_manager" not in st.session_state:
    st.session_state.context_manager = ContextManager()
if "custom_constraints" not in st.session_state:
    st.session_state.custom_constraints = {}

# Initialize handlers
audio_handler = AudioHandler()
llm_interface = LLMInterface()
slide_processor = SlideProcessor()

def clean_markdown_for_tts(text: str) -> str:
    """
    Clean markdown formatting for TTS processing.
    
    Parameters
    ----------
    text : str
        The markdown text to clean.
        
    Returns
    -------
    str
        Cleaned text suitable for TTS.
    """
    # Remove bold/italic markers
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*(.+?)\*', r'\1', text)      # Italic
    text = re.sub(r'_(.+?)_', r'\1', text)        # Alternate italic
    text = re.sub(r'__(.+?)__', r'\1', text)      # Alternate bold
    
    # Remove markdown links, keeping only the text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    
    # Remove code markers
    text = re.sub(r'`(.+?)`', r'\1', text)
    text = re.sub(r'```[\s\S]*?```', '', text)    # Remove code blocks
    
    # Remove list markers
    text = re.sub(r'^\s*[-*+]\s+', '', text, flags=re.MULTILINE)  # Unordered lists
    text = re.sub(r'^\s*\d+\.\s+', '', text, flags=re.MULTILINE)  # Ordered lists
    
    # Remove headers
    text = re.sub(r'#+\s+', '', text)
    
    # Remove any remaining asterisks
    text = text.replace('*', '')
    
    # Clean up extra whitespace
    text = re.sub(r'\s+', ' ', text)
    text = text.strip()
    
    return text

def get_cache_status_color(current_idx: int, cols_per_row: int) -> str:
    """
    Get the color for the cache status indicator.
    
    Parameters
    ----------
    current_idx : int
        Current slide index.
    cols_per_row : int
        Number of slides per row.
        
    Returns
    -------
    str
        Color code for the status indicator.
    """
    next_idx = current_idx + cols_per_row
    
    if next_idx >= len(st.session_state.slides):
        return "gray"  # No more slides to cache
    elif st.session_state.is_caching and st.session_state.next_cache_index == next_idx:
        return "yellow"  # Currently caching next page
    
    # Check if next page is in cache
    next_page_key = get_page_key(next_idx, cols_per_row)
    cached_data = st.session_state.page_cache.get(next_page_key)
    if cached_data and cached_data["start_idx"] == next_idx:
        return "green"  # Next page is cached
    
    return "red"  # Next page not cached yet

def clear_chat():
    """Clear the chat history."""
    st.session_state.chat_history = []

def get_page_key(start_idx: int, cols_per_row: int) -> str:
    """
    Generate a unique key for page caching.
    
    Parameters
    ----------
    start_idx : int
        Starting index of the page.
    cols_per_row : int
        Number of slides per row.
        
    Returns
    -------
    str
        Unique page key.
    """
    # Ensure we use consistent key format regardless of actual end index
    return f"page_{start_idx}"

def cache_page_data(
    start_idx: int,
    cols_per_row: int,
    grid_image,
    explanation: str,
    audio_file: Path,
    model_name: str,
    tts_provider: str
) -> None:
    """
    Cache all data related to a page.
    
    Parameters
    ----------
    start_idx : int
        Starting index of the page.
    cols_per_row : int
        Number of slides per row.
    grid_image : Image
        Grid image of the slides.
    explanation : str
        Generated explanation text.
    audio_file : Path
        Path to the generated audio file.
    model_name : str
        Name of the model used for generation.
    tts_provider : str
        TTS provider used.
    """
    page_key = get_page_key(start_idx, cols_per_row)
    st.session_state.page_cache[page_key] = {
        "start_idx": start_idx,
        "grid_image": grid_image,
        "explanation": explanation,
        "audio_file": audio_file,
        "model_name": model_name,
        "tts_provider": tts_provider,
        "timestamp": pd.Timestamp.now(),
        "saved_to_notes": start_idx in st.session_state.saved_notes
    }

def get_cached_page(start_idx: int, cols_per_row: int) -> dict:
    """
    Get cached page data if available.
    
    Parameters
    ----------
    start_idx : int
        Starting index of the page.
    cols_per_row : int
        Number of slides per row.
        
    Returns
    -------
    dict
        Cached page data or None if not found.
    """
    page_key = get_page_key(start_idx, cols_per_row)
    cached_data = st.session_state.page_cache.get(page_key)
    
    # Verify that the cached data matches the current configuration
    if cached_data and cached_data["start_idx"] == start_idx:
        return cached_data
    return None

async def preload_next_slides(
    start_idx: int,
    cols_per_row: int,
    slides: list,
    llm_interface: LLMInterface,
    audio_handler: AudioHandler,
    api_choice: str,
    gemini_model: str,
    tts_provider: str,
    kokoro_voice: str,
    kokoro_type: str
) -> None:
    """
    Preload explanations and audio for the next set of slides.
    
    Parameters
    ----------
    start_idx : int
        Starting index for the next set of slides.
    cols_per_row : int
        Number of slides per row.
    slides : list
        List of all slides.
    llm_interface : LLMInterface
        LLM interface instance.
    audio_handler : AudioHandler
        Audio handler instance.
    api_choice : str
        Selected AI provider.
    gemini_model : str
        Selected Gemini model.
    tts_provider : str
        Selected TTS provider.
    kokoro_voice : str
        Selected Kokoro voice.
    kokoro_type : str
        Selected Kokoro type.
    """
    if start_idx >= len(slides):
        return
        
    st.session_state.is_caching = True
    st.session_state.next_cache_index = start_idx
    
    try:
        # Check if page is already cached
        if get_cached_page(start_idx, cols_per_row):
            return
            
        end_idx = min(start_idx + cols_per_row, len(slides))
        next_slides = slides[start_idx:end_idx]
        
        grid_image = slide_processor.create_grid(next_slides, cols_per_row)
        explanation = await llm_interface.generate_explanation(
            grid_image,
            """Explain these slides in detail, focusing on key concepts and relationships. 
            Use the slides as context. be concise and to the point.
            your response will be used to generate a voiceover for the slides,
            start directly with the explanation, do NOT start with 'Here's a concise explanation for a voiceover' 
            or anything like that.""",
            provider=api_choice,
            model_name=gemini_model
        )
        
        # Clean markdown before TTS
        tts_text = clean_markdown_for_tts(explanation)
        
        # Convert explanation to speech
        audio_file = audio_handler.text_to_speech(
            tts_text,
            provider=tts_provider,
            voice_id=kokoro_voice if tts_provider == "Kokoro" else None,
            kokoro_type=kokoro_type.lower()
        )
        
        # Cache all page data
        cache_page_data(
            start_idx,
            cols_per_row,
            grid_image,
            explanation,
            audio_file,
            gemini_model,
            tts_provider
        )
    finally:
        st.session_state.is_caching = False
        st.session_state.next_cache_index = None

def clear_all_caches():
    """Clear all caches and reset session state for new slides."""
    st.session_state.current_slide_index = 0
    st.session_state.chat_history = []
    st.session_state.cached_responses = {}
    st.session_state.page_cache = {}
    st.session_state.is_caching = False
    st.session_state.next_cache_index = None
    st.session_state.saved_notes = set()
    st.session_state.notes_file = None
    st.session_state.context_manager.clear_context()  # Clear context

def get_notes_filename(uploaded_file_name: str, format_type: str) -> str:
    """
    Generate the notes filename from the uploaded file name.
    
    Parameters
    ----------
    uploaded_file_name : str
        The name of the uploaded file.
    format_type : str
        The desired format type ('txt' or 'docx').
    
    Returns
    -------
    str
        The generated filename with appropriate extension.
    """
    base_name = os.path.splitext(uploaded_file_name)[0]
    extension = ".txt" if format_type == "txt" else ".docx"
    return f"notes_{base_name}{extension}"

def format_markdown_for_docx(doc, text: str) -> None:
    """
    Format markdown text for Word document with proper styling using markdown library.
    
    Parameters
    ----------
    doc : Document
        The Word document to format.
    text : str
        The markdown text to format.
    """
    # Convert markdown to HTML
    html = markdown.markdown(text, extensions=['extra'])
    
    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    
    def process_element(element, parent_paragraph=None):
        """Recursively process HTML elements and apply formatting."""
        if element.name is None:
            # Text node
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
            parent_paragraph.add_run(element.string or '')
            return parent_paragraph
            
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Headers
            level = int(element.name[1])
            p = doc.add_paragraph(element.get_text())
            p.style = f'Heading {level}'
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            return p
            
        elif element.name == 'p':
            # Paragraphs
            p = doc.add_paragraph()
            for child in element.children:
                process_element(child, p)
            return p
            
        elif element.name == 'code':
            # Code blocks
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
            run = parent_paragraph.add_run(element.get_text())
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(199, 37, 78)  # Red color for code
            run.font.size = Pt(10)
            return parent_paragraph
            
        elif element.name == 'strong':
            # Bold text
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
            run = parent_paragraph.add_run(element.get_text())
            run.bold = True
            return parent_paragraph
            
        elif element.name == 'em':
            # Italic text
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
            run = parent_paragraph.add_run(element.get_text())
            run.italic = True
            return parent_paragraph
            
        elif element.name == 'ul':
            # Unordered lists
            for li in element.find_all('li', recursive=False):
                p = doc.add_paragraph(style='List Bullet')
                process_element(li, p)
            return None
            
        elif element.name == 'ol':
            # Ordered lists
            for i, li in enumerate(element.find_all('li', recursive=False), 1):
                p = doc.add_paragraph(style='List Number')
                process_element(li, p)
            return None
            
        elif element.name == 'pre':
            # Code blocks
            p = doc.add_paragraph()
            run = p.add_run(element.get_text())
            run.font.name = 'Courier New'
            run.font.color.rgb = RGBColor(199, 37, 78)
            run.font.size = Pt(10)
            return p
            
        else:
            # Process other elements recursively
            if parent_paragraph is None:
                parent_paragraph = doc.add_paragraph()
            for child in element.children:
                process_element(child, parent_paragraph)
            return parent_paragraph
    
    # Process the HTML tree
    for element in soup.children:
        process_element(element)

def save_explanation_to_notes(
    explanation: str,
    start_idx: int,
    grid_image=None
) -> None:
    """
    Save the explanation to the notes file and update the cache.
    
    Parameters
    ----------
    explanation : str
        The explanation text to save.
    start_idx : int
        The starting index of the slides being explained.
    grid_image : Optional[Image.Image]
        The grid image to include in DOCX format.
    """
    if not st.session_state.notes_file:
        return
        
    # Save the edited version if it exists, otherwise use the original
    text_to_save = st.session_state.edited_explanations.get(start_idx, explanation)
    
    # Get the format type from session state
    format_type = st.session_state.notes_format
    
    if format_type == "txt":
        with open(st.session_state.notes_file, "a", encoding="utf-8") as f:
            f.write(f"\n{text_to_save}\n----\n")
    else:  # docx
        # Create a new document if it doesn't exist
        if not os.path.exists(st.session_state.notes_file):
            doc = Document()
        else:
            doc = Document(st.session_state.notes_file)
        
        # Add the grid image if provided
        if grid_image:
            # Save image to bytes
            img_byte_arr = io.BytesIO()
            grid_image.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            
            # Add image to document
            doc.add_picture(img_byte_arr, width=Inches(6))
        
        # Format the markdown text with proper styling
        format_markdown_for_docx(doc, text_to_save)
        
        # Add separator
        doc.add_paragraph("----")
        
        # Save the document
        doc.save(st.session_state.notes_file)
    
    st.session_state.saved_notes.add(start_idx)

def clear_notes():
    """Clear the notes file and reset the saved notes cache."""
    if st.session_state.notes_file and os.path.exists(st.session_state.notes_file):
        os.remove(st.session_state.notes_file)
    st.session_state.saved_notes = set()
    st.session_state.notes_file = None

async def main():
    """Main application function."""
    # Sidebar configuration
    with st.sidebar:
        # Title with cache status indicator
        title_col1, title_col2 = st.columns([4, 1])
        with title_col1:
            st.title("⚙️ Configuration")
        with title_col2:
            cache_color = get_cache_status_color(
                st.session_state.current_slide_index,
                cols_per_row if 'cols_per_row' in locals() else 2
            )
            st.markdown(
                f"""
                <div style="
                    width: 15px;
                    height: 15px;
                    border-radius: 50%;
                    background-color: {cache_color};
                    margin-top: 15px;
                    box-shadow: 0 0 5px {cache_color};
                    display: inline-block;
                "></div>
                """,
                unsafe_allow_html=True
            )
            if st.session_state.is_caching:
                st.caption("Caching...")
            elif cache_color == "green":
                st.caption("Ready")
            elif cache_color == "gray":
                st.caption("End")
            else:
                st.caption("Waiting")

        # File upload
        uploaded_file = st.file_uploader(
            "Upload Presentation/Images",
            type=["pptx", "png", "jpg", "jpeg", "pdf"],
            help="Upload a PowerPoint file, PDF, or images"
        )
        
        # Display settings
        st.subheader("Display Settings")
        cols_per_row = st.selectbox(
            "Slides per row",
            options=[1, 2, 4],
            index=1
        )
        
        # API Selection
        st.subheader("AI Provider")
        api_choice = st.radio(
            "Select AI Provider",
            options=["Gemini", "OpenRouter"],
            index=0
        )
        
        # Gemini Model Selection
        if api_choice == "Gemini":
            gemini_model = st.selectbox(
                "Select Gemini Model",
                options=list(llm_interface.GEMINI_MODELS.keys()),
                index=0,
                help="Choose between faster (Flash Lite) or more thoughtful (Thinking) model"
            )
        else:
            gemini_model = "Gemini 2.0 Flash Lite"  # Default model
        
        # TTS Settings
        st.subheader("Text-to-Speech")
        tts_provider = st.radio(
            "Select TTS Provider",
            options=["gTTS", "Kokoro"],
            index=0
        )
        
        # Kokoro-specific settings
        if tts_provider == "Kokoro":
            kokoro_type = st.radio(
                "Kokoro TTS Type",
                options=["API", "Local"],
                index=0,
                help="Choose between web-hosted API or local model"
            )
            
            kokoro_voice = st.selectbox(
                "Kokoro Voice",
                options=["af", "ar", "cs", "de", "el", "en", "es", "fi", "fr", 
                        "hi", "hu", "it", "ja", "ko", "nl", "pl", "pt", "ru", 
                        "sv", "tr", "uk", "vi", "zh"],
                index=5,  # Default to "en"
                help="Select the voice for Kokoro TTS"
            )
        else:
            kokoro_type = "API"
            kokoro_voice = "en"
            
        # Notes section in sidebar
        st.markdown("---")
        st.subheader("📝 Notes")
        
        # Format selection in sidebar
        if st.session_state.notes_file is None:
            format_type = st.selectbox(
                "Select notes format:",
                options=["docx", "txt"],
                index=0 if st.session_state.notes_format == "docx" else 1,
                help="Choose the format for saving notes"
            )
            st.session_state.notes_format = format_type
            st.session_state.notes_file = get_notes_filename(uploaded_file.name, format_type)
        else:
            # Display current format
            st.info(f"Current format: {st.session_state.notes_format.upper()}")
        
        # Download notes button
        if st.session_state.notes_file and os.path.exists(st.session_state.notes_file):
            if st.session_state.notes_format == "txt":
                with open(st.session_state.notes_file, "r", encoding="utf-8") as f:
                    notes_content = f.read()
                st.download_button(
                    "📥 Download Notes",
                    notes_content,
                    file_name=os.path.basename(st.session_state.notes_file),
                    mime="text/plain",
                    help="Download the saved notes as a text file"
                )
            else:  # docx
                with open(st.session_state.notes_file, "rb") as f:
                    notes_content = f.read()
                st.download_button(
                    "📥 Download Notes",
                    notes_content,
                    file_name=os.path.basename(st.session_state.notes_file),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    help="Download the saved notes as a Word document"
                )
            
            # Clear notes button
            if st.button("🗑️ Clear Notes", help="Delete all saved notes"):
                clear_notes()
                st.rerun()
            
            # Option to change format
            if st.button("📝 Change Format", help="Change the notes format (will clear existing notes)"):
                clear_notes()
                st.rerun()

        # Add Context Settings section
        st.markdown("---")
        st.subheader("🔄 Context Settings")
        
        # Context size slider
        context_size = st.slider(
            "Number of slides in local context",
            min_value=1,
            max_value=5,
            value=3,
            help="Maximum number of previous slides to keep in immediate context"
        )
        if context_size != st.session_state.context_manager.context_size:
            st.session_state.context_manager.context_size = context_size
        
        # Constraint Management
        st.markdown("### 📋 Prompt Constraints")
        
        # Predefined constraints
        st.markdown("#### Predefined Constraints")
        for category, constraints in DEFAULT_CONSTRAINTS.items():
            with st.expander(f"{category.title()} Constraints"):
                for c_id, c_text in constraints.items():
                    if st.checkbox(
                        c_text,
                        key=f"constraint_{c_id}",
                        value=c_id in st.session_state.context_manager.active_constraints
                    ):
                        st.session_state.context_manager.add_constraint(c_id, c_text)
                        st.session_state.context_manager.activate_constraint(c_id)
                    else:
                        st.session_state.context_manager.deactivate_constraint(c_id)
        
        # Custom constraints
        st.markdown("#### Custom Constraints")
        with st.expander("Add Custom Constraint"):
            custom_text = st.text_area(
                "Enter custom constraint:",
                key="custom_constraint_input",
                help="Describe your custom constraint for the explanation"
            )
            if st.button("Add Custom Constraint"):
                if custom_text.strip():
                    constraint_id = f"custom_{len(st.session_state.custom_constraints)}"
                    st.session_state.custom_constraints[constraint_id] = custom_text
                    st.session_state.context_manager.add_constraint(constraint_id, custom_text)
                    st.session_state.context_manager.activate_constraint(constraint_id)
                    st.success("Custom constraint added!")
        
        # Display active custom constraints
        if st.session_state.custom_constraints:
            st.markdown("#### Active Custom Constraints")
            for c_id, c_text in st.session_state.custom_constraints.items():
                if st.checkbox(
                    c_text,
                    key=f"custom_{c_id}",
                    value=c_id in st.session_state.context_manager.active_constraints
                ):
                    st.session_state.context_manager.activate_constraint(c_id)
                else:
                    st.session_state.context_manager.deactivate_constraint(c_id)

    # Process uploaded file
    if uploaded_file is not None:
        # Clear all caches if this is a new file
        if not st.session_state.slides or uploaded_file.name != getattr(st.session_state, 'last_uploaded_file', None):
            with st.spinner("Processing new slides..."):
                clear_all_caches()
                st.session_state.slides = slide_processor.process_file(uploaded_file)
                st.session_state.last_uploaded_file = uploaded_file.name
                st.session_state.notes_file = None  # Reset notes file to trigger format selection

    # Main content area
    if not st.session_state.slides:
        st.title("🎯 Slide Explainer")
        st.markdown("""
        Welcome to Slide Explainer! Get started by:
        1. Uploading your presentation or images
        2. Configuring your preferred settings
        3. Exploring your content with AI-powered explanations
        """)
        return

    # Layout for slides and chat
    col1, col2 = st.columns([3, 1])
    
    with col1:
        # Slide counter and title in the same row
        header_col1, header_col2 = st.columns([3, 1])
        with header_col1:
            st.subheader("📑 Slides")
            # Add progress bar
            total_slides = len(st.session_state.slides)
            progress = (st.session_state.current_slide_index + 1) / total_slides
            st.progress(progress, "Slide Progress")
        with header_col2:
            total_pages = (total_slides + cols_per_row - 1) // cols_per_row
            current_page = (st.session_state.current_slide_index // cols_per_row) + 1
            st.markdown(
                f"""
                <div style="
                    background-color: rgba(255, 255, 255, 0.1);
                    padding: 5px 10px;
                    border-radius: 10px;
                    text-align: center;
                    margin-top: 5px;
                ">
                    Page {current_page} of {total_pages}
                    <br>
                    <small style="color: gray;">
                        ({st.session_state.current_slide_index + 1}-{min(st.session_state.current_slide_index + cols_per_row, total_slides)} of {total_slides} slides)
                    </small>
                </div>
                """,
                unsafe_allow_html=True
            )
        
        # Display current slides
        start_idx = st.session_state.current_slide_index
        end_idx = min(start_idx + cols_per_row, len(st.session_state.slides))
        current_slides = st.session_state.slides[start_idx:end_idx]
        
        if current_slides:
            # Check cache first
            cached_page = get_cached_page(start_idx, cols_per_row)
            
            if cached_page:
                # Use cached data
                grid_image = cached_page["grid_image"]
                st.image(grid_image, use_container_width=True)
                
                # Move audio player right under the image
                st.audio(str(cached_page["audio_file"]))
                
                # Navigation controls right after audio player
                nav_container = st.container()
                with nav_container:
                    nav_col1, nav_col2, nav_col3 = st.columns([2, 1, 2])
                    with nav_col1:
                        prev_clicked = st.button(
                            "⬅️ Previous",
                            use_container_width=True,
                            key="prev_button"
                        )
                    with nav_col2:
                        # Save notes button
                        is_saved = start_idx in st.session_state.saved_notes
                        save_clicked = st.button(
                            "✓ Saved" if is_saved else "📝 Save Notes",
                            help="Save explanation to notes" if not is_saved else "Already saved to notes",
                            key=f"save_notes_{start_idx}",
                            use_container_width=True
                        )
                    with nav_col3:
                        next_clicked = st.button(
                            "Next ➡️",
                            use_container_width=True,
                            key="next_button"
                        )
                
                # Display explanation text after controls in an editable text area
                edited_text = st.text_area(
                    "Edit explanation before saving:",
                    value=cached_page["explanation"],
                    height=300,
                    key=f"edit_explanation_{start_idx}"
                )
                
                # Store edited text in session state
                st.session_state.edited_explanations[start_idx] = edited_text
                
                # Handle save action
                if save_clicked and not is_saved:
                    save_explanation_to_notes(
                        edited_text,
                        start_idx,
                        grid_image if st.session_state.notes_format == "docx" else None
                    )
                    cache_page_data(
                        start_idx,
                        cols_per_row,
                        cached_page["grid_image"],
                        edited_text,  # Use edited text in cache
                        cached_page["audio_file"],
                        cached_page["model_name"],
                        cached_page["tts_provider"]
                    )
                    st.rerun()
                
                # Handle navigation
                if prev_clicked and st.session_state.current_slide_index > 0:
                    st.session_state.current_slide_index -= cols_per_row
                    st.rerun()
                elif next_clicked and end_idx < len(st.session_state.slides):
                    st.session_state.current_slide_index += cols_per_row
                    st.rerun()
                
                # Display explanation text
                st.markdown(cached_page["explanation"])
                
            else:
                # Generate new data
                grid_image = slide_processor.create_grid(current_slides, cols_per_row)
                st.image(grid_image, use_container_width=True)
                
                with st.spinner("Generating explanation..."):
                    # Build prompt using context manager
                    # Instead of trying to access text attribute, we'll describe the images
                    current_slides_desc = f"A set of {len(current_slides)} slides being displayed in a grid format. "
                    current_slides_desc += "These slides are part of the presentation and need to be explained in detail. "
                    current_slides_desc += f"Currently showing slides {start_idx + 1} to {end_idx} of the presentation."
                    
                    prompt = st.session_state.context_manager.build_prompt(current_slides_desc)
                    
                    explanation = await llm_interface.generate_explanation(
                        grid_image,
                        prompt,
                        provider=api_choice,
                        model_name=gemini_model
                    )
                    
                    # Add explanation to context
                    st.session_state.context_manager.add_explanation(
                        explanation,
                        st.session_state.current_slide_index,
                        metadata={
                            'model': gemini_model,
                            'provider': api_choice
                        }
                    )
                    
                    # Clean markdown before TTS
                    tts_text = clean_markdown_for_tts(explanation)
                    
                    # Convert explanation to speech
                    audio_file = audio_handler.text_to_speech(
                        tts_text,
                        provider=tts_provider,
                        voice_id=kokoro_voice if tts_provider == "Kokoro" else None,
                        kokoro_type=kokoro_type.lower()
                    )
                    
                    # Cache the page data
                    cache_page_data(
                        start_idx,
                        cols_per_row,
                        grid_image,
                        explanation,
                        audio_file,
                        gemini_model,
                        tts_provider
                    )
                    
                    # Display audio player and controls in the same order as cached version
                    st.audio(str(audio_file))
                    
                    # Navigation controls
                    nav_container = st.container()
                    with nav_container:
                        nav_col1, nav_col2, nav_col3 = st.columns([2, 1, 2])
                        with nav_col1:
                            prev_clicked = st.button(
                                "⬅️ Previous",
                                use_container_width=True,
                                key="prev_button"
                            )
                        with nav_col2:
                            # Save notes button
                            is_saved = start_idx in st.session_state.saved_notes
                            save_clicked = st.button(
                                "✓ Saved" if is_saved else "📝 Save Notes",
                                help="Save explanation to notes" if not is_saved else "Already saved to notes",
                                key=f"save_notes_{start_idx}",
                                use_container_width=True
                            )
                        with nav_col3:
                            next_clicked = st.button(
                                "Next ➡️",
                                use_container_width=True,
                                key="next_button"
                            )
                    
                    # Display explanation text in editable text area
                    edited_text = st.text_area(
                        "Edit explanation before saving:",
                        value=explanation,
                        height=300,
                        key=f"edit_explanation_{start_idx}"
                    )
                    
                    # Store edited text in session state
                    st.session_state.edited_explanations[start_idx] = edited_text
            
            # Preload next slides in background
            next_start_idx = start_idx + cols_per_row
            asyncio.create_task(preload_next_slides(
                next_start_idx,
                cols_per_row,
                st.session_state.slides,
                llm_interface,
                audio_handler,
                api_choice,
                gemini_model,
                tts_provider,
                kokoro_voice,
                kokoro_type
            ))
    
    with col2:
        # Chat header with clear button
        chat_col1, chat_col2 = st.columns([3, 1])
        with chat_col1:
            st.subheader("💭 Chat")
        with chat_col2:
            if st.button("🗑️", help="Clear chat history"):
                clear_chat()
        
        # Display chat history
        for message in st.session_state.chat_history:
            with st.chat_message(message["role"]):
                st.write(message["content"])
        
        # Chat input
        if prompt := st.chat_input("Ask about the slides..."):
            st.session_state.chat_history.append({
                "role": "user",
                "content": prompt
            })
            
            # Generate response
            with st.spinner("Thinking..."):
                response = await llm_interface.chat_response(
                    st.session_state.chat_history,
                    provider=api_choice,
                    model_name=gemini_model
                )
                
                st.session_state.chat_history.append({
                    "role": "assistant",
                    "content": response
                })

if __name__ == "__main__":
    asyncio.run(main()) 